import streamlit as st
import tempfile
import os
import io
import numpy as np
import time
import threading
import queue
import subprocess
import traceback
import logging
from datetime import datetime
import math

# Optional dependencies with graceful fallbacks
try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    import speech_recognition as sr
    SPEECH_RECOGNITION_AVAILABLE = True
except ImportError:
    SPEECH_RECOGNITION_AVAILABLE = False

try:
    import pyttsx3
    TTS_AVAILABLE = True
except ImportError:
    TTS_AVAILABLE = False

try:
    from scipy.io.wavfile import write, read
    SCIPY_AVAILABLE = True
except ImportError:
    SCIPY_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    from docx import Document
    DOCUMENT_SUPPORT = True
except ImportError:
    DOCUMENT_SUPPORT = False

try:
    from sentence_transformers import SentenceTransformer
    import faiss
    EMBEDDINGS_SUPPORT = True
except ImportError:
    EMBEDDINGS_SUPPORT = False

try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    ollama = None
    OLLAMA_AVAILABLE = False

# ==========================
# BROWSER MICROPHONE PERMISSION REQUEST
# ==========================
def request_microphone_permission():
    """JavaScript to automatically request microphone permission"""
    permission_js = """
    <div style="text-align: center; padding: 20px; border: 2px solid #007bff; border-radius: 10px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; margin: 15px 0;">
        <h3>🎤 Microphone Setup</h3>
        <p>Click below to enable voice recording</p>
        <button onclick="requestMicPermission()" style="background: #28a745; color: white; border: none; padding: 10px 20px; border-radius: 5px; cursor: pointer; font-size: 16px;">
            🔓 Enable Microphone
        </button>
        <div id="mic-status" style="margin-top: 15px; font-weight: bold;"></div>
    </div>

    <script>
    function requestMicPermission() {
        const statusDiv = document.getElementById('mic-status');
        statusDiv.innerHTML = '🔄 Requesting permission...';
        
        if (navigator.mediaDevices && navigator.mediaDevices.getUserMedia) {
            navigator.mediaDevices.getUserMedia({audio: true})
                .then(function(stream) {
                    // Permission granted
                    stream.getTracks().forEach(track => track.stop());
                    statusDiv.innerHTML = '✅ Microphone enabled! You can now use voice recording below.';
                    statusDiv.style.color = '#90EE90';
                })
                .catch(function(error) {
                    // Permission denied
                    statusDiv.innerHTML = `
                        ❌ Access denied! 
                        <br><small>1. Click 🎤 icon in address bar → Allow<br>
                        2. Or go to chrome://settings/content/microphone</small>
                    `;
                    statusDiv.style.color = '#FFB6C1';
                });
        } else {
            statusDiv.innerHTML = '❌ Browser not supported';
            statusDiv.style.color = '#FFB6C1';
        }
    }
    </script>
    """
    
    st.components.v1.html(permission_js, height=180)

# ==========================
# AUDIO ANALYSIS FUNCTIONS
# ==========================
def freq_to_note(frequency):
    """Convert frequency to musical note"""
    if frequency <= 0:
        return "N/A"
    
    NOTE_NAMES = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']
    A4 = 440.0
    
    try:
        semitones = 12 * math.log2(frequency / A4)
        note_index = int(round(semitones)) % 12
        octave = 4 + (int(round(semitones)) + 9) // 12
        return f"{NOTE_NAMES[note_index]}{octave}"
    except:
        return "N/A"

def analyze_audio_frequency(audio_data, sample_rate=44100):
    """Comprehensive audio analysis"""
    if not SCIPY_AVAILABLE:
        return {
            'duration': 0, 'sample_rate': sample_rate, 'max_amplitude': 0,
            'rms_level': 0, 'peak_frequency': 0, 'audio_quality': 'Install scipy for analysis',
            'dominant_frequencies': [], 'frequency_spectrum': {'frequencies': [], 'amplitudes': []}
        }
    
    try:
        # Convert audio data
        if hasattr(audio_data, 'getvalue'):
            temp_file = io.BytesIO(audio_data.getvalue())
            sample_rate, audio_array = read(temp_file)
        else:
            audio_array = np.array(audio_data)
        
        # Ensure mono
        if len(audio_array.shape) > 1:
            audio_array = audio_array[:, 0]
        
        # Normalize
        if audio_array.dtype != np.float32:
            if audio_array.dtype == np.int16:
                audio_array = audio_array.astype(np.float32) / 32768.0
            else:
                audio_array = audio_array.astype(np.float32)
        
        # FFT Analysis
        fft_result = np.abs(np.fft.rfft(audio_array))
        frequencies = np.fft.rfftfreq(len(audio_array), d=1.0/sample_rate)
        
        # Find dominant frequencies
        peak_indices = np.argsort(fft_result)[-10:][::-1]
        dominant_freqs = frequencies[peak_indices]
        dominant_amps = fft_result[peak_indices]
        
        # Filter valid frequencies (>20 Hz)
        valid_mask = dominant_freqs > 20
        dominant_freqs = dominant_freqs[valid_mask][:5]
        dominant_amps = dominant_amps[valid_mask][:5]
        
        # Audio statistics
        max_amplitude = np.max(np.abs(audio_array))
        rms_level = np.sqrt(np.mean(audio_array**2))
        duration = len(audio_array) / sample_rate
        
        # Peak frequency
        valid_freq_mask = frequencies > 20
        if np.any(valid_freq_mask):
            valid_fft = fft_result[valid_freq_mask]
            valid_frequencies = frequencies[valid_freq_mask]
            peak_frequency = valid_frequencies[np.argmax(valid_fft)]
        else:
            peak_frequency = 0
        
        # Quality assessment
        if max_amplitude > 0.1:
            quality = "Excellent"
        elif max_amplitude > 0.05:
            quality = "Good"  
        elif max_amplitude > 0.01:
            quality = "Fair"
        elif max_amplitude > 0.001:
            quality = "Poor"
        else:
            quality = "Very Poor"
        
        return {
            'duration': duration,
            'sample_rate': sample_rate,
            'max_amplitude': max_amplitude,
            'rms_level': rms_level,
            'peak_frequency': peak_frequency,
            'dominant_frequencies': dominant_freqs,
            'dominant_amplitudes': dominant_amps,
            'frequency_spectrum': {'frequencies': frequencies, 'amplitudes': fft_result},
            'audio_quality': quality
        }
        
    except Exception as e:
        st.error(f"Audio analysis failed: {e}")
        return None

def create_frequency_visualization(analysis_data):
    """Create Plotly frequency visualization"""
    if not analysis_data or not PLOTLY_AVAILABLE:
        return None
    
    try:
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=('Frequency Spectrum', 'Dominant Frequencies', 'Audio Level', 'Peak Analysis'),
            specs=[[{"type": "xy"}, {"type": "xy"}], [{"type": "indicator"}, {"type": "xy"}]]
        )
        
        # 1. Frequency spectrum
        freqs = analysis_data['frequency_spectrum']['frequencies']
        amps = analysis_data['frequency_spectrum']['amplitudes']
        audible_mask = (freqs >= 20) & (freqs <= 8000)
        
        fig.add_trace(
            go.Scatter(x=freqs[audible_mask], y=amps[audible_mask], mode='lines', 
                      name='Spectrum', line=dict(color='cyan', width=2)),
            row=1, col=1
        )
        
        # 2. Dominant frequencies
        dom_freqs = analysis_data.get('dominant_frequencies', [])
        if len(dom_freqs) > 0:
            freq_labels = [f"{f:.0f}Hz\n({freq_to_note(f)})" for f in dom_freqs]
            fig.add_trace(
                go.Bar(x=freq_labels, y=analysis_data.get('dominant_amplitudes', []), 
                      marker_color='orange', name='Dominant'),
                row=1, col=2
            )
        
        # 3. Audio level gauge
        fig.add_trace(
            go.Indicator(
                mode="gauge+number", value=analysis_data['max_amplitude'],
                title={'text': "Audio Level"},
                gauge={'axis': {'range': [None, 1.0]},
                       'bar': {'color': "lightgreen"},
                       'steps': [{'range': [0, 0.001], 'color': "red"},
                                {'range': [0.001, 0.01], 'color': "yellow"},
                                {'range': [0.01, 1.0], 'color': "green"}]}
            ),
            row=2, col=1
        )
        
        # 4. Peak frequency analysis
        peak_freq = analysis_data['peak_frequency']
        fig.add_trace(
            go.Bar(x=['Peak Frequency'], y=[peak_freq], 
                  text=[f"{peak_freq:.0f} Hz\n{freq_to_note(peak_freq)}"],
                  textposition='auto', marker_color='red'),
            row=2, col=2
        )
        
        fig.update_layout(height=600, title="🎵 Audio Analysis Dashboard", showlegend=False)
        return fig
        
    except Exception as e:
        st.error(f"Visualization failed: {e}")
        return None

# ==========================
# DOCUMENT PROCESSING
# ==========================
def extract_pdf(file_bytes: bytes, name: str):
    if not DOCUMENT_SUPPORT:
        raise ImportError("Install PyMuPDF: pip install PyMuPDF")
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append({"doc": name, "page": i+1, "text": text.strip()})
        doc.close()
        return pages
    except Exception as e:
        raise Exception(f"PDF extraction failed: {e}")

def extract_docx(file_bytes: bytes, name: str):
    if not DOCUMENT_SUPPORT:
        raise ImportError("Install python-docx: pip install python-docx")
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_text.append(row_text)
        
        text = "\n".join(paragraphs + table_text)
        return [{"doc": name, "page": 1, "text": text}] if text.strip() else []
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {e}")

def extract_txt(file_bytes: bytes, name: str):
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return [{"doc": name, "page": 1, "text": text.strip()}]
        except UnicodeDecodeError:
            continue
    raise Exception(f"Could not decode {name}")

def extract_any(uploaded_file):
    name = uploaded_file.name
    data = uploaded_file.read()
    
    if name.lower().endswith('.pdf'):
        return extract_pdf(data, name)
    elif name.lower().endswith('.docx'):
        return extract_docx(data, name)
    elif name.lower().endswith('.txt'):
        return extract_txt(data, name)
    else:
        raise ValueError(f"Unsupported file: {name}")

def chunk_text(pages, target_chars=1200, overlap_chars=200):
    chunks = []
    for p in pages:
        text = p["text"].strip()
        if not text:
            continue
        
        # Smart sentence chunking
        sentences = text.replace('!', '.').replace('?', '.').split('.')
        sentences = [s.strip() for s in sentences if s.strip()]
        
        current_chunk = ""
        for sentence in sentences:
            test_chunk = current_chunk + ". " + sentence if current_chunk else sentence
            
            if len(test_chunk) <= target_chars:
                current_chunk = test_chunk
            else:
                if len(current_chunk) >= 100:
                    chunks.append({"doc": p["doc"], "page": p["page"], "text": current_chunk.strip()})
                
                if overlap_chars > 0 and len(current_chunk) > overlap_chars:
                    overlap_text = current_chunk[-overlap_chars:]
                    current_chunk = overlap_text + ". " + sentence
                else:
                    current_chunk = sentence
        
        if len(current_chunk) >= 100:
            chunks.append({"doc": p["doc"], "page": p["page"], "text": current_chunk.strip()})
    
    return chunks

# ==========================
# RAG SYSTEM
# ==========================
@st.cache_resource(show_spinner=False)
def load_embedder():
    if not EMBEDDINGS_SUPPORT:
        return None
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

def l2_normalize(embeddings):
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True) + 1e-12
    return embeddings / norms

class VectorIndex:
    def __init__(self, dim):
        if not EMBEDDINGS_SUPPORT:
            return
        self.index = faiss.IndexFlatIP(dim)
        self.metadata = []
    
    def add(self, embeddings, metadata):
        if not EMBEDDINGS_SUPPORT:
            return
        self.index.add(embeddings.astype(np.float32))
        self.metadata.extend(metadata)
    
    def search(self, query_vec, top_k=5):
        if not EMBEDDINGS_SUPPORT or self.index.ntotal == 0:
            return []
        
        top_k = min(top_k, self.index.ntotal)
        scores, indices = self.index.search(query_vec.astype(np.float32), top_k)
        
        results = []
        for idx, score in zip(indices[0], scores[0]):
            if idx != -1:
                results.append((self.metadata[idx], float(score)))
        return results

# ==========================
# LLM INTEGRATION
# ==========================
def get_ollama_models():
    if not OLLAMA_AVAILABLE:
        return ["llama3.2:3b"]
    
    try:
        result = subprocess.run(["ollama", "list"], capture_output=True, text=True, timeout=10)
        if result.returncode != 0:
            return ["llama3.2:3b"]
        
        models = []
        for line in result.stdout.strip().split('\n')[1:]:
            if line.strip():
                models.append(line.split()[0])
        return models if models else ["llama3.2:3b"]
    except:
        return ["llama3.2:3b"]

def call_llm_ollama(model_name, prompt, temperature=0.2, max_tokens=512):
    if not OLLAMA_AVAILABLE:
        raise RuntimeError("Ollama not available")
    
    response = ollama.chat(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        options={"temperature": temperature, "num_predict": max_tokens}
    )
    return response["message"]["content"]

def build_prompt(query, retrieved_chunks, max_context_chars=4000):
    ctx_parts = []
    char_budget = max_context_chars
    
    for i, chunk in enumerate(retrieved_chunks, 1):
        piece = f"[SOURCE {i}: {chunk['doc']} | Page {chunk['page']}]\n{chunk['text'].strip()}\n" + "-"*50 + "\n"
        
        if len(piece) <= char_budget:
            ctx_parts.append(piece)
            char_budget -= len(piece)
        else:
            break
    
    context = "".join(ctx_parts)
    
    system = f"""You are an expert AI assistant. Answer the user's question using ONLY the provided sources.

INSTRUCTIONS:
- Use ONLY the information from the provided sources
- If no answer exists in sources, say "I cannot find this information in the provided documents"
- Cite sources using [SOURCE X] format
- Be thorough but concise

AVAILABLE SOURCES: {len(retrieved_chunks)} documents"""

    return f"""{system}

=== SOURCE DOCUMENTS ===
{context}

=== USER QUESTION ===
{query}

=== YOUR ANSWER ===
"""

# ==========================
# MAIN STREAMLIT APP
# ==========================
def main():
    st.set_page_config(page_title="🧠 Smart Assistant", layout="wide")
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem; border-radius: 15px; color: white; text-align: center; margin-bottom: 2rem;
    }
    .feature-card {
        border: 1px solid #e0e0e0; border-radius: 10px; padding: 1.5rem; margin: 1rem 0; background: #f8f9fa;
    }
    .status-success { color: #28a745; font-weight: bold; }
    .status-error { color: #dc3545; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🧠 Smart Document Assistant</h1>
        <p>Advanced Voice-Enabled AI Document Chat System</p>
        <p><em>Real-time Audio Analysis • Speech Recognition • Document QA • Vector Search</em></p>
    </div>
    """, unsafe_allow_html=True)

    # Browser microphone permission
    st.subheader("🎤 Step 1: Enable Microphone Access")
    request_microphone_permission()

    # Initialize session state
    for key, default in [
        ("doc_chunks", []), ("embedder", None), ("vec_index", None), 
        ("embeddings", None), ("audio_analysis", None), ("transcript", ""), 
        ("chat_history", []), ("debug_mode", False)
    ]:
        if key not in st.session_state:
            if key == "embedder" and EMBEDDINGS_SUPPORT:
                st.session_state[key] = load_embedder()
            else:
                st.session_state[key] = default

    # Sidebar Configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # System status
        st.subheader("📊 System Status")
        
        components = [
            ("🎤 Voice Recording", True, "Built-in"),
            ("📊 Audio Analysis", SCIPY_AVAILABLE, "scipy"),
            ("📈 Visualizations", PLOTLY_AVAILABLE, "plotly"),
            ("🗣️ Speech-to-Text", SPEECH_RECOGNITION_AVAILABLE, "SpeechRecognition"),
            ("🔊 Text-to-Speech", TTS_AVAILABLE, "pyttsx3"),
            ("📚 Documents", DOCUMENT_SUPPORT, "PyMuPDF, python-docx"),
            ("🔍 Vector Search", EMBEDDINGS_SUPPORT, "sentence-transformers, faiss"),
            ("🤖 LLM", OLLAMA_AVAILABLE, "ollama")
        ]
        
        for name, available, pkg in components:
            status = "status-success" if available else "status-error"
            icon = "✅" if available else "❌"
            st.markdown(f'<span class="{status}">{icon} {name}</span>', unsafe_allow_html=True)
        
        # Model settings
        if OLLAMA_AVAILABLE:
            st.subheader("🤖 AI Settings")
            models = get_ollama_models()
            model_choice = st.selectbox("Model", models, index=0)
            temperature = st.slider("Temperature", 0.0, 1.0, 0.2, 0.05)
            max_tokens = st.slider("Max Tokens", 128, 2048, 512, 64)
        else:
            model_choice, temperature, max_tokens = "llama3.2:3b", 0.2, 512
        
        # RAG settings
        if EMBEDDINGS_SUPPORT:
            st.subheader("🔍 Search Settings")
            top_k = st.slider("Top-K Documents", 1, 15, 5)
        else:
            top_k = 5
        
        # Document upload
        if DOCUMENT_SUPPORT:
            st.subheader("📚 Document Upload")
            uploaded_files = st.file_uploader(
                "Upload Files", type=["pdf", "docx", "txt"], accept_multiple_files=True
            )
            
            chunk_size = st.slider("Chunk Size", 600, 2000, 1200, 100)
            chunk_overlap = st.slider("Overlap", 50, 400, 200, 25)
            
            if st.button("📚 Build Index") and uploaded_files:
                with st.spinner("Building knowledge index..."):
                    all_chunks = []
                    
                    for file in uploaded_files:
                        try:
                            pages = extract_any(file)
                            chunks = chunk_text(pages, chunk_size, chunk_overlap)
                            all_chunks.extend(chunks)
                            st.success(f"✅ {file.name}: {len(chunks)} chunks")
                        except Exception as e:
                            st.error(f"❌ {file.name}: {e}")
                    
                    if all_chunks and EMBEDDINGS_SUPPORT and st.session_state.embedder:
                        texts = [chunk["text"] for chunk in all_chunks]
                        embeddings = st.session_state.embedder.encode(texts, convert_to_numpy=True)
                        embeddings = l2_normalize(embeddings)
                        
                        vec_index = VectorIndex(embeddings.shape[1])
                        vec_index.add(embeddings, all_chunks)
                        
                        st.session_state.doc_chunks = all_chunks
                        st.session_state.vec_index = vec_index
                        st.session_state.embeddings = embeddings
                        
                        st.balloons()
                        st.success(f"🎉 Index built: {len(all_chunks):,} chunks!")

    # Main interface
    tab1, tab2, tab3, tab4 = st.tabs(["🎙️ Voice Interface", "💬 Document Chat", "📊 Analysis", "🛠️ System"])

    # Voice Interface Tab
    with tab1:
        st.header("🎙️ Voice Recording & Analysis")
        
        st.markdown("""
        <div class="feature-card">
        <h4>📋 Instructions:</h4>
        <ol>
        <li><strong>Enable microphone</strong> using the button above ↑</li>
        <li><strong>Click record</strong> below and speak clearly</li>
        <li><strong>Get instant transcription</strong> and frequency analysis</li>
        <li><strong>Use transcript</strong> for document questions</li>
        </ol>
        </div>
        """, unsafe_allow_html=True)
        
        # Voice recording using st.audio_input (proven to work)
        st.subheader("🎤 Record Your Voice")
        audio_bytes = st.audio_input("🔴 Click to record your message")
        
        if audio_bytes is not None:
            st.success("✅ **Recording successful!**")
            
            # Display audio
            st.audio(audio_bytes, format='audio/wav')
            
            # Audio info
            audio_size = len(audio_bytes.getvalue())
            st.info(f"📊 **Audio:** {audio_size:,} bytes • WAV format")
            
            # Audio analysis
            if SCIPY_AVAILABLE:
                with st.spinner("🔬 Analyzing frequency spectrum..."):
                    analysis = analyze_audio_frequency(audio_bytes)
                    st.session_state.audio_analysis = analysis
                    
                    if analysis:
                        # Quick stats
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("⏱️ Duration", f"{analysis['duration']:.1f}s")
                        with col2:
                            st.metric("🎯 Quality", analysis['audio_quality'])
                        with col3:
                            st.metric("📊 Peak", f"{analysis['peak_frequency']:.0f} Hz")
                        with col4:
                            st.metric("🎵 Note", freq_to_note(analysis['peak_frequency']))
            else:
                st.warning("⚠️ Install scipy for audio analysis: `pip install scipy`")
            
            # Speech-to-text
            st.subheader("🗣️ Speech Recognition")
            
            col_stt1, col_stt2 = st.columns([1, 2])
            
            with col_stt1:
                if st.button("📝 **Transcribe Speech**", type="primary"):
                    if SPEECH_RECOGNITION_AVAILABLE:
                        with st.spinner("🔄 Converting speech to text..."):
                            try:
                                temp_file = io.BytesIO(audio_bytes.getvalue())
                                recognizer = sr.Recognizer()
                                recognizer.energy_threshold = 300
                                recognizer.dynamic_energy_threshold = True
                                
                                with sr.AudioFile(temp_file) as source:
                                    recognizer.adjust_for_ambient_noise(source)
                                    audio_data = recognizer.record(source)
                                    transcript = recognizer.recognize_google(audio_data, language='en-US')
                                    st.session_state.transcript = transcript
                                    
                            except sr.UnknownValueError:
                                st.error("❌ Speech unclear - try speaking more clearly")
                            except Exception as e:
                                st.error(f"❌ Transcription failed: {e}")
                    else:
                        st.error("❌ Install SpeechRecognition: `pip install SpeechRecognition`")
            
            with col_stt2:
                if st.session_state.transcript:
                    st.success("✅ **Transcription Complete:**")
                    st.text_area("Your Speech:", value=st.session_state.transcript, height=80)
            
            # Action buttons
            st.subheader("💾 Actions")
            col_a1, col_a2, col_a3 = st.columns(3)
            
            with col_a1:
                st.download_button(
                    "📥 Download Audio",
                    data=audio_bytes.getvalue(),
                    file_name=f"recording_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav",
                    mime="audio/wav"
                )
            
            with col_a2:
                if st.session_state.transcript:
                    st.download_button(
                        "📄 Download Transcript",
                        data=st.session_state.transcript,
                        file_name=f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                    )
            
            with col_a3:
                if st.button("🗑️ Clear All"):
                    st.session_state.transcript = ""
                    st.session_state.audio_analysis = None
                    st.rerun()

    # Document Chat Tab
    with tab2:
        st.header("💬 AI-Powered Document Chat")
        
        if not st.session_state.vec_index:
            st.warning("⚠️ **No documents loaded!** Upload files in the sidebar first.")
            
            if DOCUMENT_SUPPORT and EMBEDDINGS_SUPPORT:
                st.info("📋 **Setup:** Upload documents → Build Index → Ask questions")
            else:
                missing = []
                if not DOCUMENT_SUPPORT:
                    missing.append("Document processing")
                if not EMBEDDINGS_SUPPORT:
                    missing.append("Vector search")
                st.error(f"❌ **Missing:** {', '.join(missing)}")
        else:
            st.success(f"✅ **Ready:** {len(st.session_state.doc_chunks):,} chunks loaded")
            
            # Question input
            st.subheader("❓ Ask Your Question")
            
            text_query = st.text_input(
                "Type question:", 
                placeholder="Ask anything about your documents...",
                key="chat_input"
            )
            
            # Voice input option
            if st.session_state.transcript:
                st.info(f"🎤 **Voice question:** '{st.session_state.transcript}'")
                if st.button("Use Voice Question"):
                    text_query = st.session_state.transcript
            
            # Process question
            if text_query:
                if OLLAMA_AVAILABLE:
                    with st.spinner("🔍 Searching documents and generating answer..."):
                        try:
                            # Search
                            query_emb = st.session_state.embedder.encode([text_query], convert_to_numpy=True)
                            query_emb = l2_normalize(query_emb)
                            results = st.session_state.vec_index.search(query_emb, top_k)
                            chunks = [r[0] for r in results]
                            
                            if chunks:
                                # Generate answer
                                prompt = build_prompt(text_query, chunks)
                                answer = call_llm_ollama(model_choice, prompt, temperature, max_tokens)
                                
                                # Display
                                st.subheader("🤖 AI Assistant Response")
                                st.markdown(answer)
                                
                                # Stats
                                col_s1, col_s2, col_s3 = st.columns(3)
                                with col_s1:
                                    st.metric("📊 Sources", len(chunks))
                                with col_s2:
                                    st.metric("📝 Response", f"{len(answer)} chars")
                                with col_s3:
                                    st.metric("🎯 Model", model_choice)
                                
                                # Sources
                                with st.expander(f"📚 Source Documents ({len(chunks)})"):
                                    for i, chunk in enumerate(chunks, 1):
                                        st.markdown(f"### 📄 Source {i}: {chunk['doc']} (Page {chunk['page']})")
                                        preview = chunk['text'][:300] + "..." if len(chunk['text']) > 300 else chunk['text']
                                        st.markdown(f"*{preview}*")
                                
                                # Actions
                                col_act1, col_act2 = st.columns(2)
                                
                                with col_act1:
                                    if st.button("🔊 Speak Answer") and TTS_AVAILABLE:
                                        with st.spinner("Converting to speech..."):
                                            try:
                                                engine = pyttsx3.init()
                                                clean_answer = answer.replace('[', '').replace(']', '').replace('*', '')
                                                engine.say(clean_answer)
                                                engine.runAndWait()
                                                st.success("✅ Speech completed")
                                            except Exception as e:
                                                st.error(f"TTS failed: {e}")
                                
                                with col_act2:
                                    if st.button("💾 Save to History"):
                                        entry = {
                                            "timestamp": datetime.now(),
                                            "question": text_query,
                                            "answer": answer,
                                            "sources": len(chunks)
                                        }
                                        st.session_state.chat_history.append(entry)
                                        st.success("✅ Saved!")
                            else:
                                st.warning("⚠️ No relevant documents found")
                                
                        except Exception as e:
                            st.error(f"❌ Error: {e}")
                else:
                    st.error("❌ Ollama not available - install from https://ollama.ai")

    # Analysis Tab
    with tab3:
        st.header("📊 Audio Analysis Dashboard")
        
        if st.session_state.audio_analysis:
            analysis = st.session_state.audio_analysis
            
            # Plotly visualization
            if PLOTLY_AVAILABLE:
                fig = create_frequency_visualization(analysis)
                if fig:
                    st.plotly_chart(fig, use_container_width=True)
            else:
                st.warning("⚠️ Install plotly: `pip install plotly`")
            
            # Detailed analysis
            st.subheader("📈 Detailed Analysis")
            
            col_d1, col_d2 = st.columns(2)
            
            with col_d1:
                st.markdown("**🎤 Voice Characteristics:**")
                freqs = analysis.get('dominant_frequencies', [])
                for i, freq in enumerate(freqs[:3], 1):
                    note = freq_to_note(freq)
                    st.write(f"• **Frequency {i}:** {freq:.0f} Hz ({note})")
            
            with col_d2:
                st.markdown("**📊 Quality Metrics:**")
                metrics = [
                    ("Duration", f"{analysis['duration']:.2f}s"),
                    ("Sample Rate", f"{analysis['sample_rate']:,} Hz"),
                    ("Max Amplitude", f"{analysis['max_amplitude']:.4f}"),
                    ("Quality", analysis['audio_quality'])
                ]
                for name, value in metrics:
                    st.write(f"• **{name}:** {value}")
        else:
            st.info("🎤 Record audio in Voice Interface tab to see analysis")

    # System Tab
    with tab4:
        st.header("🛠️ System Tools & Diagnostics")
        
        # Comprehensive status
        for name, available, pkg in components:
            with st.expander(f"{'🟢' if available else '🔴'} {name}"):
                st.write(f"**Package:** {pkg}")
                if available:
                    st.success("✅ Working correctly")
                else:
                    st.error("❌ Not available")
                    if pkg != "Built-in":
                        st.code(f"pip install {pkg}")
        
        # Session info
        st.subheader("💾 Session Information")
        
        session_info = {
            "Documents Loaded": len(st.session_state.doc_chunks),
            "Vector Index": "Ready" if st.session_state.vec_index else "Not Built",
            "Audio Analysis": "Available" if st.session_state.audio_analysis else "None",
            "Transcript": "Available" if st.session_state.transcript else "Empty",
            "Chat History": f"{len(st.session_state.chat_history)} messages"
        }
        
        for key, value in session_info.items():
            st.write(f"• **{key}:** {value}")
        
        if st.button("🗑️ Clear All Session Data"):
            for key in list(st.session_state.keys()):
                if key != "embedder":
                    del st.session_state[key]
            st.success("✅ Session cleared")
            st.rerun()

    # Footer
    st.divider()
    st.markdown("""
    ---
    **🧠 Smart Document Assistant v3.0** | *Production-Ready Voice AI System*
    
    **Features:** Voice Recording • Audio Analysis • Speech Recognition • Document QA • Vector Search • LLM Integration
    """)

if __name__ == "__main__":
    main()
